#!/usr/bin/env python3
"""
md-to-docx.py

Konvertuje set markdown fajlova u jedan .docx sa akademskim stilom:
  - Times New Roman 12pt
  - Dvostruki prored (line spacing 2.0)
  - Hijerarhijski naslovi (H1 16pt bold, H2 14pt bold, H3 12pt bold italic)
  - Tabele sa grid border-ima
  - Code blokovi u Consolas 10pt sa sivom pozadinom

Upotreba: python3 scripts/md-to-docx.py
Output:   report/kompletni-report.docx
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm

ROOT = Path(__file__).resolve().parent.parent
REPORT = ROOT / "report"

# Redosled fajlova u objedinjenom dokumentu
FILES_IN_ORDER = [
    "report.md",
    "analysis.md",
    "comparison-table.md",
    "appendix-run-logs.md",
]

# Akademski stil
BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(12)
HEADING_SIZES = {1: Pt(16), 2: Pt(14), 3: Pt(12)}
CODE_FONT = "Consolas"
CODE_SIZE = Pt(10)
CODE_BG = "F0F0F0"  # svetlo siva


def set_cell_border(cell):
    """Dodaj grid border (sve 4 strane) na ćeliju tabele."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for border_name in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def shade_cell(cell, fill="DDDDDD"):
    """Oboj pozadinu ćelije (koristi se za header redove)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_paragraph_format(paragraph, *, line_spacing=2.0, space_after=Pt(6),
                         space_before=Pt(0), alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent=None):
    """Standardni paragraf: dvostruki prored, justify, razmak posle."""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = space_after
    pf.space_before = space_before
    paragraph.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent


def add_runs_with_inline_formatting(paragraph, text, *, base_font=BODY_FONT,
                                     base_size=BODY_SIZE, base_bold=False, base_italic=False):
    """Parsiraj inline markdown (**bold**, *italic*, `code`) i dodaj kao run-ove."""
    # Podeli na tokene: **bold**, *italic*, `code`, ili običan tekst
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.font.name = base_font
            run.font.size = base_size
            run.bold = base_bold
            run.italic = base_italic
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = CODE_FONT
            run.font.size = CODE_SIZE
        elif token.startswith("["):
            # [text](url) → samo text (URL ne prikazujemo u docx)
            label = re.match(r"\[([^\]]+)\]", token).group(1)
            run = paragraph.add_run(label)
            run.italic = True
        else:
            run = paragraph.add_run(token)
        if not run.font.name:
            run.font.name = base_font
        if not run.font.size:
            run.font.size = base_size
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = base_font
        run.font.size = base_size
        run.bold = base_bold
        run.italic = base_italic


def add_heading(doc, text, level):
    """Dodaj heading sa odgovarajućom veličinom i bold-om."""
    sizes = {1: 18, 2: 14, 3: 13, 4: 12}
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12 if level == 1 else 8)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(sizes.get(level, 12))
    run.bold = True
    if level == 3:
        run.italic = True
    return p


def add_paragraph(doc, text):
    """Dodaj paragraf sa inline formatiranjem."""
    if not text.strip():
        return None
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0, space_after=Pt(6),
                         first_line_indent=Cm(0.5))
    add_runs_with_inline_formatting(p, text)
    return p


def add_code_block(doc, code):
    """Dodaj code blok u Consolas sa sivom pozadinom."""
    for line in code.splitlines():
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.left_indent = Cm(0.5)
        run = p.add_run(line if line else " ")
        run.font.name = CODE_FONT
        run.font.size = CODE_SIZE
        # Pozadina paragrafa
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), CODE_BG)
        pPr.append(shd)


def add_table(doc, lines):
    """Parsiraj markdown tabelu i dodaj kao Word tabelu sa grid borderima."""
    # Filtriranje separatora (|---|---|) i header-a
    rows = []
    for line in lines:
        s = line.strip()
        if not s or not s.startswith("|"):
            continue
        if re.match(r"^\|[\s\-:|]+\|$", s):
            continue  # separator red
        cells = [c.strip() for c in s.strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.line_spacing = 1.15
            pf.space_after = Pt(0)
            add_runs_with_inline_formatting(p, cell_text, base_size=Pt(11))
            set_cell_border(cell)
            if i == 0:
                # Header red: bold + siva pozadina
                shade_cell(cell, "DDDDDD")
                for run in p.runs:
                    run.bold = True


def parse_and_add(doc, md_text):
    """Parsiraj markdown liniju po liniju i dodaj u docx."""
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2), level)
            i += 1
            continue

        # Code block
        if s.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # preskoči zatvarajući ```
            add_code_block(doc, "\n".join(code_lines))
            continue

        # Horizontal rule
        if re.match(r"^---+$", s) or re.match(r"^\*\*\*+$", s):
            p = doc.add_paragraph("─" * 40)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Tabela
        if s.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue

        # Lista
        if re.match(r"^[\-\*]\s+", s):
            # Sakupljaj celu listu
            list_lines = []
            while i < len(lines) and re.match(r"^\s*[\-\*]\s+", lines[i]):
                list_lines.append(lines[i])
                i += 1
            for ll in list_lines:
                indent = len(ll) - len(ll.lstrip())
                content = re.sub(r"^\s*[\-\*]\s+", "", ll)
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.line_spacing = 1.5
                pf.space_after = Pt(3)
                pf.left_indent = Cm(0.5 + indent * 0.25)
                run = p.add_run("• ")
                run.font.name = BODY_FONT
                run.font.size = BODY_SIZE
                add_runs_with_inline_formatting(p, content)
            continue

        # Numerisana lista
        if re.match(r"^\d+\.\s+", s):
            list_lines = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                list_lines.append(lines[i])
                i += 1
            for n, ll in enumerate(list_lines, 1):
                content = re.sub(r"^\s*\d+\.\s+", "", ll)
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.line_spacing = 1.5
                pf.space_after = Pt(3)
                pf.left_indent = Cm(0.75)
                run = p.add_run(f"{n}. ")
                run.font.name = BODY_FONT
                run.font.size = BODY_SIZE
                add_runs_with_inline_formatting(p, content)
            continue

        # Prazan red
        if not s:
            i += 1
            continue

        # Blockquote (> ...)
        if s.startswith(">"):
            content = s.lstrip(">").strip()
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing = 1.5
            pf.left_indent = Cm(1.0)
            pf.space_after = Pt(6)
            run = p.add_run("┃ " + content)
            run.font.name = BODY_FONT
            run.font.size = BODY_SIZE
            run.italic = True
            i += 1
            continue

        # Obican paragraf (sakupljaj susedne ne-prazne linije)
        para_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith("#") or nxt.startswith("```")
                    or nxt.startswith("|") or nxt.startswith(">")
                    or re.match(r"^[\-\*]\s+", nxt) or re.match(r"^\d+\.\s+", nxt)
                    or re.match(r"^---+$", nxt)):
                break
            para_lines.append(lines[i])
            i += 1
        text = " ".join(l.strip() for l in para_lines)
        add_paragraph(doc, text)


def main():
    doc = Document()

    # Margine: 2.5cm sa svih strana (akademski standard)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default stil fonta
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE

    for fname in FILES_IN_ORDER:
        path = REPORT / fname
        if not path.exists():
            print(f"[warn] {path} ne postoji, preskačem")
            continue
        text = path.read_text(encoding="utf-8")
        # Ukloni prvi H1 (koristimo kao page title ako je potrebno, ovde samo
        # normalno parsiramo)
        parse_and_add(doc, text)
        # Page break između fajlova
        if fname != FILES_IN_ORDER[-1]:
            doc.add_page_break()

    out = REPORT / "kompletni-report.docx"
    doc.save(out)
    print(f"[ok] Sačuvano: {out} ({out.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    main()
